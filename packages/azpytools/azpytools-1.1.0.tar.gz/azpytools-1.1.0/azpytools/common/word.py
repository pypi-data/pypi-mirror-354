from docx import Document
from docx.shared import Inches, Cm,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_COLOR
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from typing import List, Optional, Union
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
from docx.shared import RGBColor
from docx.enum.style import WD_BUILTIN_STYLE
from saplib.llm import llmABAP
import re
import json
from azsaprfc.saprfc import SapServer,ConnParams
from dataclasses import asdict



class word_doc:
    DEFAULT_FONT_NAME = '微软雅黑'
    DEFAULT_FONT_SIZE = 12
    DEFAULT_FONT_BOLD = False
    DEFAULT_STYLE = WD_BUILTIN_STYLE.NORMAL
    def __init__(self,fileName:str):
        self.fileName = fileName
        self.doc = Document()

    # def create_custom_style(self, style_name: str, font_name: str = '微软雅黑', font_size: int = 12,
    #                        font_color: RGBColor = RGBColor(0, 0, 0), bold: bool = False,
    #                        alignment: WD_PARAGRAPH_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT.LEFT):
    #     """创建自定义样式
    #     Args:
    #         style_name: 样式名称
    #         font_name: 字体名称
    #         font_size: 字体大小
    #         font_color: 字体颜色
    #         bold: 是否加粗
    #         alignment: 段落对齐方式
    #     """
    #     style = self.doc.styles.add_style(style_name, 1)  # 1 表示段落样式
    #     font = style.font
    #     font.name = font_name
    #     font.size = Pt(font_size)
    #     font.color.rgb = font_color
    #     font.bold = bold
    #     style.paragraph_format.alignment = alignment
    #     return style
    def add_title(self, text: str, level: int = 0) -> None:
        """添加标题"""
        self.doc.add_heading(text, level)


    def add_image(self, image_path: str, width_cm: float = 4.8) -> None:
        """添加图片并居中"""
        try:
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(image_path, width=Cm(width_cm))
        except FileNotFoundError:
            print(f'未找到图片文件: {image_path}')

    def add_page_break(self) -> None:
        """添加分页符"""
        self.doc.add_page_break()

    def add_paragraph_markdown_text(self, markdown_text: str,  fontname:str,size:int,bold:bool, style: Optional[str] = None) -> None:
        """添加Markdown文本"""
        # 转换 Markdown 为 word 格式
        paragraphs = markdown_text.split('\n')
        for para in paragraphs:
            if para.startswith('#'):  # 标题处理
                level = para.count('#') - 1  # 获取标题级别（从0开始）
                text = re.sub(r'^#+\s*', '', para)  # 去除前面的#和空格
                self.add_paragraph(text,fontname = fontname,size = size,bold = bold , style='Heading'+str(level))
            elif para.startswith('*'):  # 列表项处理（简化处理）
                text = re.sub(r'^\*\s*', '', para).replace('***', '').replace('**', '')  # 去除前面的*和空格
                self.add_paragraph(text,fontname = fontname,size = size,bold = bold , style=style)
            else:  # 普通段落处理
                self.add_paragraph(para,fontname = fontname,size = size,bold = bold , style=style)




    def add_paragraph(self, text: str,fontname:str,size:int,bold:bool, style: Optional[str] = None) -> Paragraph:
        """添加段落"""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        font = run.font
        font.name = fontname  # 设置字体为宋体
        font.size = Pt(size)  # 设置字体大小为22磅
        font.bold = bold    # 设置加粗
        if style:
            para.style = style
        # for run in para.runs:
        #     run.font.name = fontname  # 设置字体为宋体
        #     run.font.size = Pt(size)  # 设置字体大小为22磅
        #     run.font.bold = bold    # 设置加粗
        return para

    def create_table(self, headers: List[str], data: List[List[Union[str, int, float]]]) -> Table:
        """创建表格"""
        table = self.doc.add_table(rows=1, cols=len(headers))
        # 设置表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        # 添加数据行
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        return table

    def merge_cells(self, table: Table, start_row: int, start_col: int, end_row: int, end_col: int) -> None:
        """合并表格单元格
        Args:
            table: 要操作的表格对象
            start_row: 起始行索引(从0开始)
            start_col: 起始列索引(从0开始)
            end_row: 结束行索引
            end_col: 结束列索引
        """
        cell1 = table.cell(start_row, start_col)
        cell2 = table.cell(end_row, end_col)
        cell1.merge(cell2)

    def cell_fill_color(self, table: Table, start_row: int, start_col: int, end_row: int, end_col: int, colorHex: str) -> None:
        """填充表格单元格颜色
        Args:
            table: 要操作的表格对象
            start_row: 起始行索引(从0开始)
            start_col: 起始列索引(从0开始)
            end_row: 结束行索引
            end_col: 结束列索引
            colorHex: 颜色值，例如'#FF0000'表示红色 也有用str(rgbColor(255,0,0))"""
        # 设置单元格背景颜色，这里以灰色为例，颜色代码为 'D3D3D3'
        for row in range(start_row, end_row + 1):
            for col in range(start_col, end_col + 1):
                cell = table.cell(row, col)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), colorHex)
                tcPr.append(shd)

    def add_toc(self) -> None:
        """添加自动目录"""
        # 添加目录标题
        self.add_title("目录", level=1)

        # 创建目录段落
        paragraph = self.doc.add_paragraph()
        paragraph.style = "TOC Heading"

        # 添加目录字段代码
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar)

        instrText = OxmlElement('w:instrText')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instrText)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

    def author(self,author:str) -> None:
        """添加作者"""
        self.doc.core_properties.author = author
        self.doc.core_properties.last_modified_by = author
        self.doc.core_properties.comments = ''  # 添加备注
        # self.doc.part.core_properties.comments = ''  # 添加备注
        self.doc.part.core_properties.modified = datetime.datetime.now()
        self.doc.part.core_properties.created = datetime.datetime.now()
        self.doc.part.core_properties.revision = 1
        self.doc.part.core_properties.version = 1
        self.doc.part.core_properties.language = '' #'zh-CN'
        self.doc.part.core_properties.category = '' #'文档'
        self.doc.part.core_properties.content_status = '' #'最终'
        self.doc.part.core_properties.keywords = '' #'关键字'
        self.doc.part.core_properties.subject = '' #'主题'
        self.doc.part.core_properties.title = '' #'标题'
        self.doc.part.core_properties.description = '' #'描述'
        self.doc.part.core_properties.identifier = '' #'标识符'
        # self.doc.part.core_properties.creator = '' #'创建者'
        # self.doc.part.core_properties.last_modified_by = '' #'最后修改者'


    def save(self) -> None:
        """保存文档"""
        self.doc.save(self.fileName)

    def add_header(self, text: str, font_name: str = "微软雅黑", font_size: int = 10,
                color: RGBColor = RGBColor(0, 0, 0)) -> None:
        section = self.doc.sections[0]
        """添加页眉"""
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = color

    def add_footer(self, text: str, font_name: str = "微软雅黑", font_size: int = 9,
                color: RGBColor = RGBColor(128, 128, 128), add_page_number: bool = True) -> None:
        section = self.doc.sections[0]
        """添加页脚"""
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = color

        if add_page_number:
            footer_para.add_run("\t")
            run = footer_para.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar)

            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            run._r.append(instrText)

            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar)
