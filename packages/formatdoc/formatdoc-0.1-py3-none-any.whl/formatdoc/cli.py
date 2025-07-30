from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def Add_Header(Section, text="My Document Header"):
    header = Section.header
    paragraph = header.paragraphs[0]
    paragraph.text = text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(12)
    run.font.name = 'Arial'

def Add_Footer_with_Page_Numbers(Section):
    footer = Section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def format_txt_file_to_docx(input_file, outputfile):
    document = Document()

    section = document.sections[0]
    Add_Header(section)
    Add_Footer_with_Page_Numbers(section)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    with open(input_file, 'r', encoding='utf-8') as file:
        for line in file:
            paragraph = document.add_paragraph(line.strip())
            paragraph.paragraph_format.space_after = Pt(8)
    document.save(outputfile)
    print(f"Document saved as: {outputfile}")


format_txt_file_to_docx("input.txt", "formatted_Output.docx")

def main():
    print("Formatdoc fonctionne bien !")
